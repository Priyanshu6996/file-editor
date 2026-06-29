"""
================================================================================
 Automatic Reimbursement Document Generator
================================================================================
Reads student fee data from students.xlsx and fills it into the existing
template.docx (Word) to produce Reimbursement_Output.docx.

This version is built for the real "MASOOM" reimbursement template, where
the letterhead, student table, totals, payment terms, amount-in-words line,
and signature block all live inside ONE single Word table (each "row" in
the layout is really a table row, and most static rows are six merged
cells repeating the same text). The script locates everything by searching
for the distinctive text already in the template instead of hardcoding
row/column numbers, so it keeps working even if a few rows are added or
removed above the student table.

Run with:
    python app.py

Requires (see requirements.txt):
    pandas, python-docx, openpyxl, num2words
================================================================================
"""

import os
import re
import copy

import pandas as pd
from docx import Document
from docx.table import _Row
from num2words import num2words


# --------------------------------------------------------------------------
# CONFIG: file names (everything lives next to this script)
# --------------------------------------------------------------------------
EXCEL_FILE = "students.xlsx"
TEMPLATE_FILE = "template.docx"
OUTPUT_FILE = "Reimbursement_Output.docx"

STUDENT_SHARE_PCT = 0.10   # 10% paid by student
MASOOM_SHARE_PCT = 0.90    # 90% paid by Masoom

REQUIRED_COLUMNS = ["Name", "Course", "Fees"]


# ==========================================================================
# STEP 1: LOAD AND VALIDATE THE EXCEL FILE
# ==========================================================================
def load_excel(filepath):
    """
    Reads the student Excel file and validates that it exists and has
    the required columns. Returns a pandas DataFrame, or None on error
    (error message is printed, the program does not crash).
    """
    if not os.path.exists(filepath):
        print("students.xlsx not found.")
        return None

    try:
        df = pd.read_excel(filepath)
    except Exception as exc:
        print(f"Could not read students.xlsx ({exc}).")
        return None

    # Normalize column names (strip spaces) so "Name " etc. still matches
    df.columns = [str(c).strip() for c in df.columns]

    missing = [col for col in REQUIRED_COLUMNS if col not in df.columns]
    if missing:
        print("Invalid Excel format.")
        return None

    # Drop completely empty rows and rows missing a name
    df = df.dropna(how="all")
    df = df[df["Name"].notna()]

    if df.empty:
        print("Invalid Excel format.")
        return None

    return df.reset_index(drop=True)


# ==========================================================================
# STEP 2: CALCULATE PER-STUDENT SHARES AND GRAND TOTALS
# ==========================================================================
def calculate_totals(df):
    """
    Adds 'Student Share' (10%) and 'Masoom Share' (90%) columns to the
    DataFrame, and returns (df, totals_dict) where totals_dict has the
    grand totals across all students.
    """
    df = df.copy()
    df["Fees"] = pd.to_numeric(df["Fees"], errors="coerce").fillna(0)
    df["Student Share"] = df["Fees"] * STUDENT_SHARE_PCT
    df["Masoom Share"] = df["Fees"] * MASOOM_SHARE_PCT

    totals = {
        "total_students": len(df),
        "total_fees": df["Fees"].sum(),
        "total_student_share": df["Student Share"].sum(),
        "total_masoom_share": df["Masoom Share"].sum(),
    }
    return df, totals


# ==========================================================================
# STEP 3: LOCATE THE STUDENT TABLE, SAMPLE ROW, AND TOTALS ROW
# ==========================================================================
def find_student_table(doc):
    """
    The whole template is one big table, so we find it by searching for
    the row that contains "Total Expenses" -- that uniquely identifies the
    table holding the student data (works even if there were several
    tables in the document).
    """
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text for cell in row.cells).lower()
            if "total expenses" in row_text:
                return table
    return None


def find_sample_and_total_row_index(table):
    """
    Finds the index of the "Total Expenses" row by its own text (never by
    position). The sample/template student row is the row directly above
    it -- in this template that is always the single example student row
    ("Dipali Arjune Bawiskar" in the original file).
    Returns (sample_row_index, total_row_index) or (None, None).
    """
    total_row_index = None
    for idx, row in enumerate(table.rows):
        row_text = " ".join(cell.text for cell in row.cells).lower()
        if "total expenses" in row_text:
            total_row_index = idx
            break

    if total_row_index is None or total_row_index == 0:
        return None, None

    sample_row_index = total_row_index - 1
    return sample_row_index, total_row_index


# ==========================================================================
# STEP 4: CLONE THE SAMPLE ROW (PRESERVES ALL FORMATTING)
# ==========================================================================
def duplicate_row(sample_tr):
    """
    Deep-copies the underlying XML element of the sample row. A deep copy
    of the XML preserves every bit of formatting (fonts, borders, shading,
    cell widths, merges) exactly as in the original row.
    """
    return copy.deepcopy(sample_tr)


def insert_row_before(target_tr, new_tr):
    """
    Inserts a cloned row element directly above the given target row
    element (the Total Expenses row), using lxml's addprevious so the
    new row always lands in the correct position regardless of how many
    rows have already been inserted.
    """
    target_tr.addprevious(new_tr)


# ==========================================================================
# STEP 5: FILL ONE STUDENT ROW WITH DATA (WITHOUT BREAKING FORMATTING)
# ==========================================================================
def set_cell_text(cell, text):
    """
    Writes text into a cell while preserving the existing run formatting
    (font, size, bold, color). Many cells in this template have one or
    more EMPTY paragraphs surrounding the paragraph that actually holds
    the visible value (e.g. an empty paragraph then the real value, or
    the real value then a trailing empty paragraph) -- so instead of
    always assuming the first paragraph holds the text, we pick whichever
    paragraph currently has real content and write into that one. If
    every paragraph is empty, we fall back to the last paragraph.
    """
    paragraphs = cell.paragraphs

    target_para = None
    for para in paragraphs:
        if any(run.text.strip() for run in para.runs):
            target_para = para
    if target_para is None:
        target_para = paragraphs[-1]

    if target_para.runs:
        target_para.runs[0].text = str(text)
        for run in target_para.runs[1:]:
            run.text = ""
    else:
        target_para.add_run(str(text))

    # Clear any other paragraph's text so old sample data doesn't linger
    for para in paragraphs:
        if para is not target_para:
            for run in para.runs:
                run.text = ""


def fill_student_row(table, new_tr, sr_no, name, course, fees, student_share, masoom_share):
    """
    Wraps the newly inserted <w:tr> element as a proper python-docx Row
    object and fills its 6 cells with the student's data. Numbers are
    written as plain integers (no thousands separator) to match the
    template's own student-row style (e.g. "36043"); the Total Expenses
    row uses comma formatting separately (see update_totals_row).
    """
    row = _Row(new_tr, table)
    set_cell_text(row.cells[0], sr_no)
    set_cell_text(row.cells[1], name)
    set_cell_text(row.cells[2], course)
    set_cell_text(row.cells[3], f"{int(round(fees))}")
    set_cell_text(row.cells[4], f"{int(round(student_share))}")
    set_cell_text(row.cells[5], f"{int(round(masoom_share))}")


# ==========================================================================
# STEP 6: AMOUNT IN WORDS
# ==========================================================================
def amount_in_words(amount):
    """
    Converts a numeric amount into words, Indian-numbering style
    (e.g. 45000 -> 'Forty Five Thousand Only').
    """
    rupees = int(round(amount))
    try:
        words = num2words(rupees, lang="en_IN")
    except NotImplementedError:
        words = num2words(rupees)
    words = words.replace("-", " ").replace(",", "")
    words = " ".join(w.capitalize() for w in words.split())
    return f"{words} Only"


# ==========================================================================
# STEP 7: UPDATE THE TOTALS ROW, STUDENT-COUNT LINE, 50% PAYMENT CELL,
#         AND THE "RS. IN WORDS" LINE
# ==========================================================================
def update_paragraph_text(paragraph, new_text):
    """
    Replaces the visible text of a paragraph while keeping the formatting
    of its first run (so font/size/bold of the line is preserved).
    """
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def update_totals_row(table, total_row_tr, totals):
    """
    Fills the Total Expenses row's last three cells (Total Fees,
    Total Student Share, Total Masoom Share) with the grand totals,
    using comma-formatted numbers to match the template's own totals-row
    style (e.g. "36,043"). The "Total Expenses" label cell is untouched.
    """
    total_row = _Row(total_row_tr, table)
    set_cell_text(total_row.cells[3], f"{int(round(totals['total_fees'])):,}")
    set_cell_text(total_row.cells[4], f"{int(round(totals['total_student_share'])):,}")
    set_cell_text(total_row.cells[5], f"{int(round(totals['total_masoom_share'])):,}")


def update_student_count_text(table, total_students):
    """
    Updates "(List of 01 student enclose)" with the actual student count,
    zero-padded to 2 digits to match the template's own "01" style.
    Edits only the specific run that holds the digits, so the rest of
    the sentence (and its bold formatting) is untouched.
    """
    new_num = f"{total_students:02d}"
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if re.search(r"list of\s*\d+", run.text, re.IGNORECASE):
                        run.text = re.sub(r"\d+", new_num, run.text, count=1)
                        return


def update_half_masoom_payment_cell(table, total_masoom_share):
    """
    Updates the numeric cell next to the "50% payment will be done after
    admission" label with Masoom Total / 2, comma-formatted to match the
    template's style.
    """
    half_payment = total_masoom_share / 2
    for row in table.rows:
        row_text = " ".join(c.text for c in row.cells).lower()
        if "50%" in row_text and "payment" in row_text:
            set_cell_text(row.cells[-1], f"{int(round(half_payment)):,}")
            return


def update_amount_in_words_cell(table, total_masoom_share):
    """
    Updates the "Rs. In word :..." line with the Masoom total converted
    into words. The label text before the colon is kept exactly as-is;
    only the words after the colon are replaced.
    """
    words_text = amount_in_words(total_masoom_share)
    for row in table.rows:
        for cell in row.cells:
            if "in word" in cell.text.lower():
                for para in cell.paragraphs:
                    if "in word" in para.text.lower():
                        prefix = para.text.split(":", 1)[0] + ":" if ":" in para.text else "Rs. In word :"
                        update_paragraph_text(para, prefix + words_text)
                        return
                return


def update_totals(doc, table, total_row_tr, totals):
    """
    Convenience wrapper that runs every "update the totals section" step:
    table totals row, student-count line, 50% payment cell, and the
    amount-in-words line.
    """
    update_totals_row(table, total_row_tr, totals)
    update_student_count_text(table, totals["total_students"])
    update_half_masoom_payment_cell(table, totals["total_masoom_share"])
    update_amount_in_words_cell(table, totals["total_masoom_share"])


# ==========================================================================
# STEP 8: SAVE THE FINAL DOCUMENT
# ==========================================================================
def save_document(doc, output_path):
    doc.save(output_path)


# ==========================================================================
# MAIN WORKFLOW
# ==========================================================================
def main():
    # ---- Load Excel ----
    print("Reading students.xlsx...")
    df = load_excel(EXCEL_FILE)
    if df is None:
        return  # error already printed, stop without crashing

    print(f"Found {len(df)} students")

    # ---- Load Word template ----
    if not os.path.exists(TEMPLATE_FILE):
        print("template.docx not found.")
        return

    try:
        doc = Document(TEMPLATE_FILE)
    except Exception as exc:
        print(f"Could not open template.docx ({exc}).")
        return

    # ---- Calculate totals ----
    print("Calculating totals...")
    df, totals = calculate_totals(df)

    # ---- Locate the student table, sample row, and totals row ----
    print("Updating template...")
    table = find_student_table(doc)
    if table is None:
        print("Invalid template format: 'Total Expenses' row not found.")
        return

    sample_row_idx, total_row_idx = find_sample_and_total_row_index(table)
    if sample_row_idx is None:
        print("Invalid template format: 'Total Expenses' row not found.")
        return

    sample_tr = table.rows[sample_row_idx]._tr
    total_tr = table.rows[total_row_idx]._tr

    # ---- Duplicate the sample row for every student and insert before totals ----
    print("Generating document...")
    students = df.to_dict(orient="records")  # safe access even for column names with spaces
    for sr_no, student in enumerate(students, start=1):
        new_tr = duplicate_row(sample_tr)
        insert_row_before(total_tr, new_tr)
        fill_student_row(
            table,
            new_tr,
            sr_no=sr_no,
            name=student["Name"],
            course=student["Course"],
            fees=student["Fees"],
            student_share=student["Student Share"],
            masoom_share=student["Masoom Share"],
        )

    # ---- Remove the original placeholder sample row ----
    sample_tr.getparent().remove(sample_tr)

    # ---- Update totals row, student count, 50% payment, amount in words ----
    update_totals(doc, table, total_tr, totals)

    # ---- Save ----
    save_document(doc, OUTPUT_FILE)
    print("Done!")
    print(f"Saved as\n{OUTPUT_FILE}")


if __name__ == "__main__":
    main()
